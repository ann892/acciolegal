"""
Accio Legal — Proposal & Engagement Letter Pipeline (v1 demo)

Flow:
  1. Read first-call transcript (Otter export)
  2. Claude extracts structured matter data — client, scope phases, pricing
  3. Komal reviews extraction in terminal (the "checkpoint") and confirms or corrects
  4. Claude drafts the full proposal narrative
  5. Render proposal to .docx with Accio brand styling
  6. Pause for client reply (sample acceptance email used in demo)
  7. Claude extracts any new info from the reply (signatory confirm, etc.)
  8. Fill Komal's engagement letter .docx template
  9. Save both files to output/

Replace samples/transcript.txt and samples/client_reply.txt with real Otter
exports / Gmail thread bodies in production. Otter webhook + Gmail API hooks
slot in at the read-points.
"""

from __future__ import annotations

import json
import os
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

import anthropic
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

ROOT = Path(__file__).parent
STYLE_REF_PATH = ROOT / "style" / "kalolytic_proposal.txt"
EL_TEMPLATE_PATH = ROOT / "templates" / "engagement_template.docx"
OUTPUT_DIR = ROOT / "output"

MODEL = "claude-sonnet-4-6"
ACCIO_NAVY = RGBColor(0x0E, 0x2A, 0x47)
ACCIO_ADDRESS = (
    "16, Gyanendra Kanan, Jagacha, GIP Colony, Howrah 711112, India"
)
ACCIO_EMAIL = "komal@acciolegal.com"
ACCIO_FIRM_LINE = (
    "Accio Legal, a partnership firm having its business address at "
    "16 Gyanendra Kanan, Jagacha, GIP Colony, Howrah 711112, India"
)

# Cached system block — sent on every Claude call so we get cache hits across
# the three-call sequence within the same matter.
SYSTEM_BLOCK = """You are a senior associate at Accio Legal, a Delhi/Howrah-based corporate law firm led by Komal Shah. You draft proposals and engagement letters for Indian startup and corporate clients. You write in Komal's voice: confident, strategic, precise on Indian statutory references (Companies Act 2013, Section 68, PAS-3, MGT-14, SH-7, etc.). You never invent pricing or scope — those come from the call transcript only. You never invent client facts — if a field is missing from the transcript, mark it as MISSING rather than guessing."""

# ---------------------------------------------------------------------------
# Claude calls
# ---------------------------------------------------------------------------

def get_client() -> anthropic.Anthropic:
    if not os.environ.get("ANTHROPIC_API_KEY"):
        sys.exit(
            "ERROR: ANTHROPIC_API_KEY not set. Copy .env.example to .env "
            "and fill in the key, then `export $(cat .env | xargs)`."
        )
    return anthropic.Anthropic()


def extract_matter_data(client: anthropic.Anthropic, transcript: str) -> dict:
    """Claude call #1 — pull structured scope/pricing/client data from transcript."""
    schema = """{
  "client": {
    "legal_name": "string — full registered legal name (e.g. 'Velara Bioworks Private Limited'). MISSING if not stated.",
    "entity_type": "string — Private Limited / LLP / Public Limited / etc.",
    "registered_address": "string — full registered office address with PIN. MISSING if not stated.",
    "primary_contact_name": "string — person on the call from client side",
    "primary_contact_email": "string — MISSING if not stated",
    "signatory_name": "string — who will sign the engagement letter on client side",
    "signatory_designation": "string — title/role (e.g. Founder & CEO)"
  },
  "matter_summary": "1–2 sentence plain-English summary of what the client is engaging Accio to do",
  "phases": [
    {
      "name": "string — short title (e.g. 'Phase I: Regularization of Pending Share Allotment')",
      "scope_summary": "1–3 sentence description of work in this phase, including statutory references Komal cited",
      "scope_bullets": ["bullet", "bullet"],
      "outcome_bullets": ["bullet", "bullet"],
      "tat_days": "string as quoted on the call (e.g. '3-5 working days')",
      "fee_inr": "integer — fee in INR as quoted on the call. null if not quoted."
    }
  ],
  "total_fee_inr": "integer — sum of phase fees, as quoted",
  "payment_terms": "string — exact payment terms Komal stated",
  "out_of_pocket_note": "string — what's excluded (gov fees, stamp duty, taxes)",
  "deferred_or_separate": "string — anything Komal said would be quoted separately later. empty string if none."
}"""

    user_msg = f"""Below is the transcript of Komal's first call with a prospective client. Extract the structured matter data.

CRITICAL RULES:
- Pricing AND scope come from the transcript only. Do not invent.
- If a field is not stated on the call, write "MISSING" (for strings) or null (for numbers).
- Use Komal's exact phrasing for statutory references (Section 68, PAS-3, MGT-14, etc.)
- Do not include phases Komal explicitly deferred (e.g. "we'll quote it then") in the main `phases` list — put those in `deferred_or_separate`.

Return ONLY a JSON object matching this schema (no preamble, no code fences):

{schema}

TRANSCRIPT:
{transcript}"""

    resp = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=[
            {"type": "text", "text": SYSTEM_BLOCK, "cache_control": {"type": "ephemeral"}},
        ],
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = resp.content[0].text.strip()
    # Strip code fences if present
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
    return json.loads(raw)


def apply_corrections(client: anthropic.Anthropic, data: dict, corrections: str) -> dict:
    """Claude call (optional) — merge Komal's free-text corrections into the data."""
    user_msg = f"""Komal reviewed the extracted matter data and provided corrections in plain text. Apply them and return the updated JSON object.

CURRENT DATA:
{json.dumps(data, indent=2)}

KOMAL'S CORRECTIONS (plain text):
{corrections}

Return ONLY the updated JSON object, same schema as before. No preamble."""
    resp = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=[{"type": "text", "text": SYSTEM_BLOCK, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
    return json.loads(raw)


def draft_proposal_narrative(client: anthropic.Anthropic, data: dict, style_ref: str) -> dict:
    """Claude call #2 — write the proposal in Accio's voice. Returns structured sections for .docx rendering."""
    user_msg = f"""Below is a STYLE REFERENCE (a real past Accio proposal — Kalolytic) and the STRUCTURED MATTER DATA for a NEW client. Draft the new proposal in Komal's voice, modeled on the style reference's structure but specific to this client's facts and pricing.

Return a JSON object with this shape:
{{
  "title": "Proposal for {{Client Name}}",
  "opening_paragraph": "one short paragraph framing why this matters for the client's commercial position — growth/fundraising/control/compliance angle. Reference the client's stage, sector, or context as drawn from the matter data.",
  "approach_bullets": ["3-5 bullets summarizing what Accio will do at a high level"],
  "scope_section_title": "Scope of Services (e.g. 'Comprehensive Equity Restructuring & ESOP Mandate')",
  "scope_section_intro": "one-line intro before the phases",
  "phases": [
    {{
      "name": "exact phase name from matter data",
      "framing_sentence": "one sentence Accio-voice framing of this phase",
      "scope_bullets": ["..."],
      "outcome_bullets": ["..."]
    }}
  ],
  "deliverables": ["bulleted list of concrete deliverables across all phases"],
  "cost_table": [
    {{"service": "Phase name", "tat": "x working days", "fee_inr": 8000}}
  ],
  "exclusions_note": "string — out-of-pocket / govt fees / taxes excluded language",
  "deferred_note": "string — anything quoted separately later. empty string if none.",
  "next_steps": ["confirmation over email", "engagement agreement sent across", "comms thread for data gathering", "delivery within timelines discussed"]
}}

CRITICAL:
- Use ONLY the pricing and scope from the matter data. Do not invent.
- Cite the same statutory references the matter data quotes (Section 68, PAS-3, MGT-14, etc.).
- Match the depth and tone of the style reference — confident, strategic, senior-counsel.
- For each phase, the framing_sentence + scope_bullets + outcome_bullets must mirror the style reference's structure.
- Output JSON ONLY, no preamble or code fences.

STYLE REFERENCE (Kalolytic — illustrative, do not copy pricing or scope):
{style_ref}

NEW MATTER DATA:
{json.dumps(data, indent=2)}"""
    resp = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        system=[{"type": "text", "text": SYSTEM_BLOCK, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
    return json.loads(raw)


def classify_reply(client: anthropic.Anthropic, reply: str, matter_data: dict) -> dict:
    """Claude call #3 — classify client reply + extract any updated/confirmed info."""
    user_msg = f"""The client replied to the proposal. Decide whether this is an acceptance, a negotiation (price/scope pushback), or a rejection. If it's an acceptance, also extract or confirm any new info — signatory name, payment readiness, kickoff date, or scope changes.

Return JSON:
{{
  "classification": "accept | negotiate | reject | needs_followup",
  "rationale": "one sentence why",
  "confirmed_signatory_name": "string or MISSING",
  "confirmed_signatory_email": "string or MISSING",
  "scope_adjustments": "string — any phases dropped, added, or modified vs. proposal. empty string if none.",
  "client_questions": ["any questions the client raised that need Komal's answer before EL goes out"],
  "ready_to_send_el": true_or_false
}}

JSON only, no preamble.

PROPOSAL MATTER DATA (what we sent):
{json.dumps(matter_data, indent=2)}

CLIENT REPLY:
{reply}"""
    resp = client.messages.create(
        model=MODEL,
        max_tokens=2048,
        system=[{"type": "text", "text": SYSTEM_BLOCK, "cache_control": {"type": "ephemeral"}}],
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
    return json.loads(raw)


# ---------------------------------------------------------------------------
# .docx rendering — Proposal
# ---------------------------------------------------------------------------

def _set_run_font(run, *, size=11, bold=False, color=None, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def render_proposal_docx(content: dict, matter_data: dict, out_path: Path) -> None:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(content["title"])
    _set_run_font(r, size=22, bold=True, color=ACCIO_NAVY)

    # Brand line
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ACCIO LEGAL")
    _set_run_font(r, size=11, bold=True, color=ACCIO_NAVY)
    doc.add_paragraph()

    # Opening
    p = doc.add_paragraph()
    r = p.add_run(content["opening_paragraph"])
    _set_run_font(r, size=11)

    p = doc.add_paragraph()
    r = p.add_run("This proposal outlines our end-to-end legal approach to:")
    _set_run_font(r, size=11)
    for b in content.get("approach_bullets", []):
        doc.add_paragraph(b, style="List Bullet")

    # Scope
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run(content.get("scope_section_title", "Scope of Services"))
    _set_run_font(r, size=15, bold=True, color=ACCIO_NAVY)

    p = doc.add_paragraph()
    r = p.add_run(content.get("scope_section_intro", ""))
    _set_run_font(r, size=11)

    for i, phase in enumerate(content.get("phases", []), start=1):
        ph = doc.add_paragraph()
        r = ph.add_run(phase["name"])
        _set_run_font(r, size=13, bold=True, color=ACCIO_NAVY)

        fr = doc.add_paragraph()
        r = fr.add_run(phase.get("framing_sentence", ""))
        _set_run_font(r, size=11)

        if phase.get("scope_bullets"):
            sh = doc.add_paragraph()
            r = sh.add_run("Scope of Work:")
            _set_run_font(r, size=11, bold=True)
            for b in phase["scope_bullets"]:
                doc.add_paragraph(b, style="List Bullet")

        if phase.get("outcome_bullets"):
            oh = doc.add_paragraph()
            r = oh.add_run("Outcome:")
            _set_run_font(r, size=11, bold=True)
            for b in phase["outcome_bullets"]:
                doc.add_paragraph(b, style="List Bullet")

    # Deliverables
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run("Deliverables")
    _set_run_font(r, size=15, bold=True, color=ACCIO_NAVY)
    for b in content.get("deliverables", []):
        doc.add_paragraph(b, style="List Bullet")

    # Cost table
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run("Summary of Costs")
    _set_run_font(r, size=15, bold=True, color=ACCIO_NAVY)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Service", "Duration / TAT", "Pricing (INR)"]):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        _set_run_font(r, size=11, bold=True, color=ACCIO_NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total = 0
    for row in content.get("cost_table", []):
        cells = table.add_row().cells
        cells[0].text = row.get("service", "")
        cells[1].text = row.get("tat", "")
        fee = row.get("fee_inr") or 0
        cells[2].text = f"{int(fee):,}" if fee else "TBD"
        total += int(fee) if fee else 0
    if total:
        cells = table.add_row().cells
        cells[0].text = "Total"
        cells[1].text = ""
        cells[2].text = f"{total:,}"
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_run_font(r, size=11, bold=True)

    p = doc.add_paragraph()
    r = p.add_run(content.get("exclusions_note", ""))
    _set_run_font(r, size=10)
    r.italic = True

    if content.get("deferred_note"):
        p = doc.add_paragraph()
        r = p.add_run(f"Note: {content['deferred_note']}")
        _set_run_font(r, size=10)
        r.italic = True

    # Next steps
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run("Next Steps")
    _set_run_font(r, size=15, bold=True, color=ACCIO_NAVY)
    for i, step in enumerate(content.get("next_steps", []), start=1):
        doc.add_paragraph(f"{i}. {step}")

    # Contact
    doc.add_paragraph()
    h = doc.add_paragraph()
    r = h.add_run("Contact")
    _set_run_font(r, size=12, bold=True, color=ACCIO_NAVY)
    for line in [
        "Komal Shah — Co-Founder & CEO",
        ACCIO_EMAIL,
        "+91 91671 25177  |  +91 82996 11658",
        ACCIO_ADDRESS,
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        _set_run_font(r, size=10)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# .docx rendering — Engagement Letter (fills Komal's existing template)
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph, find: str, replace: str) -> bool:
    """Replace text in a paragraph, preserving formatting where possible."""
    if find not in paragraph.text:
        return False
    # Concatenate runs, replace, then put it back into the first run and clear the rest.
    full = paragraph.text
    new = full.replace(find, replace)
    if not paragraph.runs:
        paragraph.add_run(new)
        return True
    paragraph.runs[0].text = new
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _insert_paragraph_after(paragraph, text: str, *, bold=False) -> None:
    new_p = deepcopy(paragraph._p)
    # Clear children then add a fresh run via python-docx after insertion.
    paragraph._p.addnext(new_p)
    # Wrap and reset
    from docx.text.paragraph import Paragraph
    inserted = Paragraph(new_p, paragraph._parent)
    for run in inserted.runs:
        run.text = ""
    run = inserted.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def render_engagement_letter(
    template_path: Path,
    matter_data: dict,
    reply_data: dict,
    out_path: Path,
) -> None:
    doc = Document(str(template_path))

    client_info = matter_data["client"]
    client_legal = client_info["legal_name"]
    client_addr = client_info["registered_address"]
    client_email = (
        reply_data.get("confirmed_signatory_email")
        if reply_data.get("confirmed_signatory_email") not in (None, "MISSING", "")
        else client_info.get("primary_contact_email", "MISSING")
    )
    signatory = (
        reply_data.get("confirmed_signatory_name")
        if reply_data.get("confirmed_signatory_name") not in (None, "MISSING", "")
        else client_info.get("signatory_name", "MISSING")
    )
    today = date.today().strftime("%dth %B %Y").replace("01th", "1st").replace("02th", "2nd").replace("03th", "3rd").replace("21th", "21st").replace("22th", "22nd").replace("23th", "23rd").replace("31th", "31st")

    # 1. Replace effective date
    for p in doc.paragraphs:
        _replace_in_paragraph(p, "26th August 2025", today)

    # 2. Insert client party block right after "BY AND BETWEEN"
    client_party = (
        f"{client_legal}, a {client_info.get('entity_type', 'company')} "
        f"having its registered office at {client_addr} "
        f"(herein referred to as the \"Client\" of the FIRST PART)"
    )
    for p in doc.paragraphs:
        if p.text.strip() == "BY AND BETWEEN":
            _insert_paragraph_after(p, client_party)
            break

    # 3. Fill blank Address: ___________________ in section 6 (Consultant address)
    for p in doc.paragraphs:
        if "Address: ___________________" in p.text:
            _replace_in_paragraph(
                p, "Address: ___________________", f"Address: {ACCIO_ADDRESS}"
            )

    # 4. Fill Client email in section 6 — find the "Email address:" right after "For the Client:"
    found_client_block = False
    for p in doc.paragraphs:
        if "For the Client:" in p.text:
            found_client_block = True
            continue
        if found_client_block and p.text.strip().startswith("Email address:"):
            new_text = f"Email address: {client_email}"
            _replace_in_paragraph(p, p.text.strip(), new_text)
            break

    # 5. Fill Annexure A scope table (Service / Duration/TAT / Pricing)
    annexure_table = None
    for table in doc.tables:
        if table.rows and "Service" in table.rows[0].cells[0].text and "Pricing" in table.rows[0].cells[-1].text:
            annexure_table = table
            break

    if annexure_table is not None:
        phases = matter_data.get("phases", [])
        # Template has 3 empty rows after header. Add more if needed.
        empty_rows = [r for r in annexure_table.rows[1:] if not any(c.text.strip() for c in r.cells[:3])]
        for i, phase in enumerate(phases):
            if i < len(empty_rows):
                row = empty_rows[i]
            else:
                row = annexure_table.add_row()
            row.cells[0].text = phase["name"]
            row.cells[1].text = phase.get("tat_days", "")
            fee = phase.get("fee_inr")
            row.cells[2].text = f"₹{int(fee):,}" if fee else "TBD"
            for c in row.cells[:3]:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)

    # 6. Signature block — fill the empty "Name:" paragraph (client side).
    # The template has two patterns we need to handle:
    #   Pattern A (Table 0): one paragraph holds "Name: \nAuthorised Signatory"
    #     (the line break is a soft return inside a single paragraph)
    #   Pattern B (Table 2): "Name: " is its own paragraph
    # Each signature cell has the client block (empty "Name:") followed by the
    # Accio block ("Name: Komal Shah") — we only fill the empty one.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Komal Shah" in p.text:
                        continue
                    if "Name: \nAuthorised Signatory" in p.text:
                        _replace_in_paragraph(
                            p, "Name: \nAuthorised Signatory",
                            f"Name: {signatory}\nAuthorised Signatory",
                        )
                    elif "Name:\nAuthorised Signatory" in p.text:
                        _replace_in_paragraph(
                            p, "Name:\nAuthorised Signatory",
                            f"Name: {signatory}\nAuthorised Signatory",
                        )
                    elif p.text.strip() == "Name:":
                        _replace_in_paragraph(p, "Name:", f"Name: {signatory}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# Checkpoint UI (terminal — replaced by WhatsApp/Slack ping in production)
# ---------------------------------------------------------------------------

def show_checkpoint(data: dict) -> None:
    print("\n" + "=" * 72)
    print("EXTRACTION CHECKPOINT — review before drafting")
    print("=" * 72)
    c = data["client"]
    print(f"\nCLIENT")
    print(f"  Legal name:     {c.get('legal_name')}")
    print(f"  Entity:         {c.get('entity_type')}")
    print(f"  Address:        {c.get('registered_address')}")
    print(f"  Contact:        {c.get('primary_contact_name')} <{c.get('primary_contact_email')}>")
    print(f"  Signatory:      {c.get('signatory_name')} ({c.get('signatory_designation')})")
    print(f"\nMATTER SUMMARY")
    print(f"  {data.get('matter_summary')}")
    print(f"\nSCOPE & PRICING")
    for p in data.get("phases", []):
        fee = p.get("fee_inr")
        fee_str = f"₹{int(fee):,}" if fee else "TBD"
        print(f"  • {p['name']}  —  {p.get('tat_days')}  —  {fee_str}")
        print(f"      {p.get('scope_summary')}")
    total = data.get("total_fee_inr")
    if total:
        print(f"\n  TOTAL: ₹{int(total):,}  (excludes govt fees, stamp duty, taxes)")
    print(f"\nPAYMENT TERMS")
    print(f"  {data.get('payment_terms')}")
    if data.get("deferred_or_separate"):
        print(f"\nDEFERRED / SEPARATE")
        print(f"  {data['deferred_or_separate']}")
    print("\n" + "=" * 72)


def get_confirmation() -> tuple[bool, str]:
    """Returns (confirmed, corrections). If confirmed=True, corrections is empty."""
    print("\nReview the above. Options:")
    print("  [Enter]   Confirm and draft proposal")
    print("  [type]    Free-text corrections (e.g. 'Phase II fee is 38000, not 35000')")
    print("  q         Abort")
    user = input("\n> ").strip()
    if not user:
        return True, ""
    if user.lower() == "q":
        sys.exit("Aborted by user.")
    return False, user


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def _parse_scenario_arg() -> str:
    """Read --scenario X from sys.argv. Default to velara for backwards compat."""
    for i, arg in enumerate(sys.argv):
        if arg == "--scenario" and i + 1 < len(sys.argv):
            return sys.argv[i + 1]
        if arg.startswith("--scenario="):
            return arg.split("=", 1)[1]
    return "velara"


def main() -> None:
    from scenarios import get_scenario
    from proposal_render import save_proposal_html

    mock_mode = "--mock" in sys.argv
    scenario_key = _parse_scenario_arg()
    scenario = get_scenario(scenario_key)

    print(f"\nAccio Legal — Proposal & EL Pipeline (demo)")
    print(f"Scenario: {scenario['label']}")
    if mock_mode:
        print("[MOCK MODE — using pre-baked outputs, no API calls]\n")
    else:
        print()

    transcript_path = scenario["transcript_path"]
    reply_path = scenario["reply_path"]
    mock_dir = scenario["mock_dir"]
    client_slug = scenario["client_slug"]

    transcript = transcript_path.read_text()
    style_ref = STYLE_REF_PATH.read_text()

    client = None if mock_mode else get_client()

    # 1. Extract
    print("→ Reading transcript and extracting matter data...")
    if mock_mode:
        data = json.loads((mock_dir / "01_matter_data.json").read_text())
    else:
        data = extract_matter_data(client, transcript)

    # 2. Checkpoint loop
    while True:
        show_checkpoint(data)
        confirmed, corrections = get_confirmation()
        if confirmed:
            break
        if mock_mode:
            print("\n[MOCK MODE — corrections flow not simulated; confirming as-is]")
            break
        print("\n→ Applying your corrections...")
        data = apply_corrections(client, data, corrections)

    # 3. Draft proposal
    print("\n→ Drafting proposal in Komal's voice...")
    if mock_mode:
        proposal_content = json.loads((mock_dir / "02_proposal_narrative.json").read_text())
    else:
        proposal_content = draft_proposal_narrative(client, data, style_ref)

    client_legal = data["client"]["legal_name"]

    # Render both formats: branded HTML (primary) and .docx (fallback)
    proposal_html_path = OUTPUT_DIR / f"{client_slug}_Proposal.html"
    save_proposal_html(proposal_content, proposal_html_path)
    print(f"  ✓ Proposal HTML saved: {proposal_html_path}")

    proposal_path = OUTPUT_DIR / f"{client_slug}_Proposal.docx"
    render_proposal_docx(proposal_content, data, proposal_path)
    print(f"  ✓ Proposal .docx saved: {proposal_path}")

    # Optional: create Gmail draft
    if "--gmail-draft" in sys.argv:
        from gmail_drafts import create_proposal_draft
        print("\n→ Creating Gmail draft for proposal...")
        try:
            result = create_proposal_draft(
                to_email=data["client"]["primary_contact_email"],
                to_name=data["client"]["primary_contact_name"],
                client_legal_name=client_legal,
                matter_summary=data.get("matter_summary", ""),
                attachment_path=proposal_path,
            )
            print(f"  ✓ Gmail draft created: {result['gmail_url']}")
            print(f"    Subject: {result['subject']}")
            print(f"    To:      {result['to']}")
        except Exception as e:
            print(f"  ✗ Gmail draft failed: {e}")
            print(f"    (Proposal .docx is still in {proposal_path})")

    # 4. Pause for client reply
    print("\n" + "=" * 72)
    print("Proposal is ready. Komal would now review it in Gmail Drafts and send.")
    print("=" * 72)
    print(f"\nIn this demo we'll use the sample client reply at {reply_path}.")
    input("Press Enter to simulate the client's response and draft the engagement letter...")

    reply = reply_path.read_text()

    # 5. Classify reply
    print("\n→ Reading client reply and classifying...")
    if mock_mode:
        reply_data = json.loads((mock_dir / "03_reply_classification.json").read_text())
    else:
        reply_data = classify_reply(client, reply, data)
    print(f"  Classification: {reply_data['classification']}  —  {reply_data['rationale']}")

    if reply_data["classification"] != "accept":
        print(f"\n  Reply was not a clean acceptance. Stopping before EL draft.")
        print(f"  Komal would handle this manually (negotiation/clarification).")
        return

    if not reply_data.get("ready_to_send_el"):
        print(f"\n  Client raised questions:")
        for q in reply_data.get("client_questions", []):
            print(f"    • {q}")
        print(f"  Komal would answer these before EL goes out.")
        # For demo, proceed anyway to show the EL draft
        print(f"\n  (Demo: proceeding to draft EL anyway so you can see the output.)")

    # 6. Render engagement letter
    print("\n→ Filling Komal's engagement letter template...")
    el_path = OUTPUT_DIR / f"{client_slug}_Engagement_Letter.docx"
    render_engagement_letter(EL_TEMPLATE_PATH, data, reply_data, el_path)
    print(f"  ✓ Engagement letter saved: {el_path}")

    # Optional: create Gmail draft for the EL
    if "--gmail-draft" in sys.argv:
        from gmail_drafts import create_engagement_letter_draft
        print("\n→ Creating Gmail draft for engagement letter...")
        try:
            signatory_email = (
                reply_data.get("confirmed_signatory_email")
                if reply_data.get("confirmed_signatory_email") not in (None, "MISSING", "")
                else data["client"].get("primary_contact_email")
            )
            signatory_name = (
                reply_data.get("confirmed_signatory_name")
                if reply_data.get("confirmed_signatory_name") not in (None, "MISSING", "")
                else data["client"].get("primary_contact_name")
            )
            result = create_engagement_letter_draft(
                to_email=signatory_email,
                to_name=signatory_name,
                client_legal_name=client_legal,
                attachment_path=el_path,
            )
            print(f"  ✓ Gmail draft created: {result['gmail_url']}")
            print(f"    Subject: {result['subject']}")
            print(f"    To:      {result['to']}")
        except Exception as e:
            print(f"  ✗ Gmail draft failed: {e}")
            print(f"    (Engagement letter .docx is still in {el_path})")

    print("\n" + "=" * 72)
    print("DONE. Both drafts are in the output/ folder.")
    print("=" * 72)
    print(f"\n  Proposal:           {proposal_path}")
    print(f"  Engagement letter:  {el_path}")
    print("\nIn production, both files would land as Gmail drafts in Komal's account.")


if __name__ == "__main__":
    main()
